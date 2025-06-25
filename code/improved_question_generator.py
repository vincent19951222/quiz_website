#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优化的糖尿病知识题目生成工具
"""

import json
import re
import random
from typing import List, Dict, Any
from docx import Document
import os

def extract_docx_content(file_path: str) -> str:
    """提取Word文档内容"""
    try:
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text.strip() + "\n"
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content += " | ".join(row_text) + "\n"
        
        return content
    except Exception as e:
        print(f"提取文档内容时出错: {e}")
        return ""

def parse_qa_content_improved(content: str) -> List[Dict[str, str]]:
    """改进的问答内容解析"""
    qa_pairs = []
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行和标题行
        if not line or line == "糖尿病基础知识100问":
            i += 1
            continue
        
        # 检查是否是问题（不以"答："开头的非空行）
        if not line.startswith('答：') and not line.startswith('答:'):
            question = line
            answer = ""
            
            # 查找下一行的答案
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                
                # 如果找到答案行
                if next_line.startswith('答：') or next_line.startswith('答:'):
                    # 提取答案内容（去掉"答："前缀）
                    answer_text = re.sub(r'^答[：:]\s*', '', next_line)
                    answer = answer_text
                    
                    # 继续读取后续的答案内容，直到遇到下一个问题
                    k = j + 1
                    while k < len(lines):
                        continuation_line = lines[k].strip()
                        if not continuation_line:
                            k += 1
                            continue
                        
                        # 如果是数字开头的行（可能是答案的编号部分）或者不是新问题，则继续添加到答案
                        if (continuation_line.startswith(('（', '(', '①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩')) or
                            re.match(r'^\d+[、.]', continuation_line) or
                            (not continuation_line.endswith('？') and not continuation_line.endswith('?') and 
                             not continuation_line.startswith('答：') and not continuation_line.startswith('答:'))):
                            answer += " " + continuation_line
                            k += 1
                        else:
                            # 遇到新问题，停止
                            break
                    
                    break
                else:
                    # 如果不是答案行，可能是问题的续行
                    if not next_line.endswith('？') and not next_line.endswith('?'):
                        j += 1
                        continue
                    else:
                        # 这是一个新问题，说明当前问题没有答案
                        break
                
                j += 1
            
            # 如果找到了有效的问答对
            if question and answer and len(question.strip()) > 3 and len(answer.strip()) > 5:
                qa_pairs.append({
                    "question": question.strip(),
                    "answer": answer.strip()
                })
            
            # 移动到下一个可能的问题位置
            i = j if 'j' in locals() else i + 1
        else:
            i += 1
    
    return qa_pairs

def generate_better_wrong_options(correct_answer: str, question: str, all_answers: List[str]) -> List[str]:
    """生成更好的错误选项"""
    wrong_options = []
    
    # 方法1：基于正确答案的关键词替换
    keyword_replacements = {
        # 数值相关
        '增加': '减少', '减少': '增加', '升高': '降低', '降低': '升高',
        '上升': '下降', '下降': '上升', '提高': '降低', '增强': '减弱',
        
        # 行为相关
        '不能': '可以', '禁止': '允许', '避免': '推荐', '预防': '治疗',
        '应该': '不应该', '必须': '不必', '需要': '不需要',
        
        # 时间相关
        '早期': '晚期', '急性': '慢性', '空腹': '餐后', '餐前': '餐后',
        '睡前': '起床后', '晨起': '睡前',
        
        # 程度相关
        '轻度': '重度', '局部': '全身', '少量': '大量', '适量': '过量',
        '正常': '异常', '健康': '患病', '有效': '无效', '安全': '危险',
        
        # 胰岛素相关
        '胰岛素': '胰高血糖素', '降血糖': '升血糖', '低血糖': '高血糖',
        
        # 数值修改
        '3.9': '2.9', '6.1': '7.1', '7.8': '8.8', '11.1': '10.1'
    }
    
    # 生成基于关键词替换的错误选项
    for original, replacement in keyword_replacements.items():
        if original in correct_answer and len(wrong_options) < 2:
            wrong_option = correct_answer.replace(original, replacement)
            if wrong_option != correct_answer and wrong_option not in wrong_options:
                wrong_options.append(wrong_option)
    
    # 方法2：从其他答案中选择相似但错误的选项
    if len(wrong_options) < 3:
        for other_answer in random.sample(all_answers, min(20, len(all_answers))):
            if other_answer != correct_answer and len(wrong_options) < 3:
                # 如果其他答案与当前问题主题相关，可能作为错误选项
                if any(word in other_answer for word in ['糖尿病', '血糖', '胰岛素'] if word in question):
                    # 截取部分内容作为错误选项
                    if len(other_answer) > 50:
                        wrong_option = other_answer[:50] + "..."
                    else:
                        wrong_option = other_answer
                    
                    if wrong_option not in wrong_options and wrong_option != correct_answer:
                        wrong_options.append(wrong_option)
    
    # 方法3：生成通用错误选项
    generic_options = [
        "以上说法都不正确",
        "需要根据个人具体情况而定",
        "医学上尚无明确定论",
        "取决于患者的具体病情",
        "需要进一步检查才能确定"
    ]
    
    for option in generic_options:
        if len(wrong_options) < 3 and option not in wrong_options:
            wrong_options.append(option)
    
    return wrong_options[:3]

def generate_multiple_choice_questions_improved(qa_pairs: List[Dict[str, str]], num_questions: int = 25) -> List[Dict[str, Any]]:
    """生成改进的单选题"""
    if len(qa_pairs) < num_questions:
        print(f"警告：只有 {len(qa_pairs)} 个问答对，将生成全部题目")
        selected_pairs = qa_pairs
    else:
        selected_pairs = random.sample(qa_pairs, num_questions)
    
    questions = []
    all_answers = [qa['answer'] for qa in qa_pairs]
    
    for i, qa in enumerate(selected_pairs, 1):
        question_text = qa['question']
        correct_answer = qa['answer']
        
        # 如果答案太长，截取前150个字符
        if len(correct_answer) > 150:
            correct_answer = correct_answer[:150] + "..."
        
        # 生成错误选项
        wrong_options = generate_better_wrong_options(correct_answer, question_text, all_answers)
        
        # 组合所有选项
        all_options = [correct_answer] + wrong_options
        random.shuffle(all_options)
        
        # 找到正确答案的索引
        correct_index = all_options.index(correct_answer)
        
        # 生成更详细的解析
        explanation = f"正确答案解析：{qa['answer']}"
        if len(explanation) > 200:
            explanation = explanation[:200] + "..."
        
        question_data = {
            "id": i,
            "question": question_text,
            "options": all_options,
            "correct_answer": correct_index,
            "explanation": explanation
        }
        
        questions.append(question_data)
    
    return questions

def main():
    """主函数"""
    doc_path = "/workspace/user_input_files/3.4糖尿病基础知识培训100问.docx"
    
    if not os.path.exists(doc_path):
        print(f"文档文件不存在: {doc_path}")
        return
    
    print("正在提取文档内容...")
    content = extract_docx_content(doc_path)
    
    if not content:
        print("无法提取文档内容")
        return
    
    print(f"文档内容长度: {len(content)} 字符")
    
    print("正在解析问答内容...")
    qa_pairs = parse_qa_content_improved(content)
    
    print(f"成功解析得到 {len(qa_pairs)} 个问答对")
    
    # 显示前几个问答对用于验证
    print("\n=== 解析结果预览 ===")
    for i, qa in enumerate(qa_pairs[:5]):
        print(f"\n问答对 {i+1}:")
        print(f"问题: {qa['question']}")
        print(f"答案: {qa['answer'][:100]}...")
    
    if len(qa_pairs) < 10:
        print("警告：解析的问答对数量较少")
        return
    
    print("\n正在生成单选题...")
    target_questions = min(25, len(qa_pairs))
    questions = generate_multiple_choice_questions_improved(qa_pairs, target_questions)
    
    # 保存生成的题目
    output_data = {
        "title": "糖尿病基础知识测试",
        "description": "基于糖尿病基础知识培训100问生成的测试题目",
        "time_limit": 30,
        "total_questions": len(questions),
        "questions": questions
    }
    
    with open('/workspace/data/quiz_questions.json', 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print(f"成功生成 {len(questions)} 道单选题")
    print("题目数据已保存到 /workspace/data/quiz_questions.json")
    
    # 显示题目预览
    print("\n=== 题目预览 ===")
    for i, q in enumerate(questions[:3]):
        print(f"\n第 {q['id']} 题: {q['question']}")
        for j, option in enumerate(q['options']):
            marker = "★" if j == q['correct_answer'] else " "
            print(f"  {chr(65+j)}. {option[:80]}{'...' if len(option) > 80 else ''} {marker}")
        print(f"解析: {q['explanation'][:100]}...")

if __name__ == "__main__":
    # 确保数据目录存在
    os.makedirs('/workspace/data', exist_ok=True)
    main()
