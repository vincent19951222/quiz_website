#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用知识库题目生成工具
支持多种文档格式，快速生成答题网站题目数据
"""

import json
import re
import random
import os
import argparse
from typing import List, Dict, Any, Optional
from docx import Document

def extract_docx_content(file_path: str) -> str:
    """提取Word文档内容"""
    try:
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text.strip() + "\n"
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content += " | ".join(row_text) + "\n"
        
        return content
    except Exception as e:
        print(f"提取Word文档内容时出错: {e}")
        return ""

def extract_txt_content(file_path: str) -> str:
    """提取文本文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"提取文本文件内容时出错: {e}")
        return ""

def extract_pdf_content(file_path: str) -> str:
    """提取PDF文档内容"""
    try:
        import PyPDF2
        content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        return content
    except ImportError:
        print("需要安装PyPDF2库来处理PDF文件: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"提取PDF文档内容时出错: {e}")
        return ""

def detect_qa_format(content: str) -> str:
    """检测问答格式类型"""
    patterns = {
        'chinese_format': r'答[：:]',  # 中文"答："格式
        'qa_format': r'^Q[：:].+A[：:]',  # Q: A: 格式
        'numbered_format': r'^\d+[、.]',  # 数字编号格式
        'simple_format': r'[？?]\s*\n.+',  # 问号后换行格式
    }
    
    for format_name, pattern in patterns.items():
        if re.search(pattern, content, re.MULTILINE):
            return format_name
    
    return 'unknown'

def parse_qa_content_universal(content: str) -> List[Dict[str, str]]:
    """通用问答内容解析"""
    qa_pairs = []
    format_type = detect_qa_format(content)
    
    print(f"检测到的文档格式: {format_type}")
    
    if format_type == 'chinese_format':
        # 中文"答："格式
        qa_pairs = parse_chinese_format(content)
    elif format_type == 'qa_format':
        # Q: A: 格式
        qa_pairs = parse_qa_format(content)
    elif format_type == 'numbered_format':
        # 数字编号格式
        qa_pairs = parse_numbered_format(content)
    else:
        # 尝试通用解析
        qa_pairs = parse_generic_format(content)
    
    return qa_pairs

def parse_chinese_format(content: str) -> List[Dict[str, str]]:
    """解析中文"答："格式"""
    qa_pairs = []
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line or any(skip in line for skip in ['目录', '索引', '第一章', '第二章']):
            i += 1
            continue
        
        # 检查是否是问题（不以"答："开头的非空行）
        if not line.startswith(('答：', '答:', '解答：', '解答:')):
            question = line
            answer = ""
            
            # 查找答案
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                
                # 找到答案行
                if next_line.startswith(('答：', '答:', '解答：', '解答:')):
                    answer_text = re.sub(r'^(答|解答)[：:]\s*', '', next_line)
                    answer = answer_text
                    
                    # 读取后续答案内容
                    k = j + 1
                    while k < len(lines):
                        cont_line = lines[k].strip()
                        if not cont_line:
                            k += 1
                            continue
                        
                        if (cont_line.startswith(('（', '(', '①', '②', '③', '④', '⑤')) or
                            re.match(r'^\d+[、.]', cont_line) or
                            (not cont_line.endswith(('？', '?')) and 
                             not cont_line.startswith(('答：', '答:')))):
                            answer += " " + cont_line
                            k += 1
                        else:
                            break
                    break
                j += 1
            
            if question and answer and len(question.strip()) > 3 and len(answer.strip()) > 5:
                qa_pairs.append({"question": question.strip(), "answer": answer.strip()})
            
            i = j if 'j' in locals() else i + 1
        else:
            i += 1
    
    return qa_pairs

def parse_qa_format(content: str) -> List[Dict[str, str]]:
    """解析Q: A:格式"""
    qa_pairs = []
    qa_blocks = re.findall(r'Q[：:]\s*(.+?)\s*A[：:]\s*(.+?)(?=Q[：:]|$)', content, re.DOTALL | re.IGNORECASE)
    
    for question, answer in qa_blocks:
        question = question.strip()
        answer = answer.strip()
        if len(question) > 3 and len(answer) > 5:
            qa_pairs.append({"question": question, "answer": answer})
    
    return qa_pairs

def parse_numbered_format(content: str) -> List[Dict[str, str]]:
    """解析数字编号格式"""
    qa_pairs = []
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 匹配数字开头的问题
        if re.match(r'^\d+[、.]\s*(.+)', line):
            question = re.sub(r'^\d+[、.]\s*', '', line)
            answer = ""
            
            # 收集后续的答案内容
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                
                # 如果遇到下一个数字编号，停止
                if re.match(r'^\d+[、.]\s*', next_line):
                    break
                
                answer += " " + next_line
                j += 1
            
            if question and answer and len(question) > 3 and len(answer) > 5:
                qa_pairs.append({"question": question, "answer": answer.strip()})
            
            i = j
        else:
            i += 1
    
    return qa_pairs

def parse_generic_format(content: str) -> List[Dict[str, str]]:
    """通用格式解析"""
    qa_pairs = []
    
    # 尝试按问号分割
    sections = re.split(r'[？?]', content)
    
    for i in range(len(sections) - 1):
        question = sections[i].strip().split('\n')[-1]  # 取最后一行作为问题
        answer_section = sections[i + 1].strip()
        
        # 提取答案（取前几行，直到遇到下一个可能的问题）
        answer_lines = answer_section.split('\n')
        answer = ""
        for line in answer_lines:
            line = line.strip()
            if line and not line.endswith(('？', '?')):
                answer += " " + line
            else:
                break
        
        if len(question) > 3 and len(answer) > 10:
            qa_pairs.append({"question": question + "？", "answer": answer.strip()})
    
    return qa_pairs

def generate_enhanced_wrong_options(correct_answer: str, question: str, all_answers: List[str], domain: str = "") -> List[str]:
    """生成增强的错误选项"""
    wrong_options = []
    
    # 领域特定的关键词替换
    domain_replacements = {
        'medical': {
            '增加': '减少', '减少': '增加', '升高': '降低', '降低': '升高',
            '不能': '可以', '禁止': '允许', '预防': '治疗', '急性': '慢性',
            '正常': '异常', '健康': '患病', '有效': '无效', '安全': '危险'
        },
        'technical': {
            '启动': '关闭', '开启': '禁用', '增加': '减少', '提高': '降低',
            '安装': '卸载', '连接': '断开', '启用': '禁用', '创建': '删除'
        },
        'business': {
            '增长': '下降', '盈利': '亏损', '成功': '失败', '优化': '恶化',
            '提升': '降低', '扩大': '缩小', '加强': '削弱', '改善': '恶化'
        },
        'legal': {
            '合法': '非法', '允许': '禁止', '有效': '无效', '责任': '免责',
            '义务': '权利', '强制': '自愿', '公开': '保密', '正当': '不当'
        }
    }
    
    # 使用通用替换加上领域特定替换
    replacements = {
        '是': '不是', '不是': '是', '正确': '错误', '错误': '正确',
        '应该': '不应该', '必须': '可以', '需要': '不需要', '能够': '不能',
        '重要': '不重要', '有效': '无效', '安全': '危险', '合适': '不合适'
    }
    
    if domain in domain_replacements:
        replacements.update(domain_replacements[domain])
    
    # 生成基于关键词替换的错误选项
    for original, replacement in replacements.items():
        if original in correct_answer and len(wrong_options) < 2:
            wrong_option = correct_answer.replace(original, replacement)
            if wrong_option != correct_answer and wrong_option not in wrong_options:
                wrong_options.append(wrong_option)
    
    # 从其他答案中选择相似选项
    if len(wrong_options) < 3:
        suitable_answers = [ans for ans in all_answers if ans != correct_answer and len(ans) < 200]
        for other_answer in random.sample(suitable_answers, min(10, len(suitable_answers))):
            if len(wrong_options) < 3:
                if len(other_answer) > 80:
                    wrong_option = other_answer[:80] + "..."
                else:
                    wrong_option = other_answer
                
                if wrong_option not in wrong_options:
                    wrong_options.append(wrong_option)
    
    # 添加通用错误选项
    generic_options = [
        "以上说法都不正确",
        "需要根据具体情况判断",
        "尚无明确规定",
        "因具体环境而异",
        "需要进一步确认"
    ]
    
    for option in generic_options:
        if len(wrong_options) < 3 and option not in wrong_options:
            wrong_options.append(option)
    
    return wrong_options[:3]

def generate_quiz_questions(qa_pairs: List[Dict[str, str]], config: Dict[str, Any]) -> Dict[str, Any]:
    """生成答题题目数据"""
    num_questions = min(config.get('num_questions', 25), len(qa_pairs))
    title = config.get('title', '知识测试')
    description = config.get('description', '基于知识库生成的测试题目')
    time_limit = config.get('time_limit', 30)
    domain = config.get('domain', '')
    
    if len(qa_pairs) < num_questions:
        print(f"警告：只有 {len(qa_pairs)} 个问答对，将生成全部题目")
        selected_pairs = qa_pairs
    else:
        selected_pairs = random.sample(qa_pairs, num_questions)
    
    questions = []
    all_answers = [qa['answer'] for qa in qa_pairs]
    
    for i, qa in enumerate(selected_pairs, 1):
        question_text = qa['question']
        correct_answer = qa['answer']
        
        # 限制答案长度
        if len(correct_answer) > 150:
            correct_answer = correct_answer[:150] + "..."
        
        # 生成错误选项
        wrong_options = generate_enhanced_wrong_options(correct_answer, question_text, all_answers, domain)
        
        # 组合选项并随机排列
        all_options = [correct_answer] + wrong_options
        random.shuffle(all_options)
        correct_index = all_options.index(correct_answer)
        
        # 生成解析
        explanation = f"正确答案解析：{qa['answer']}"
        if len(explanation) > 200:
            explanation = explanation[:200] + "..."
        
        question_data = {
            "id": i,
            "question": question_text,
            "options": all_options,
            "correct_answer": correct_index,
            "explanation": explanation
        }
        
        questions.append(question_data)
    
    return {
        "title": title,
        "description": description,
        "time_limit": time_limit,
        "total_questions": len(questions),
        "questions": questions
    }

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='通用知识库题目生成工具')
    parser.add_argument('input_file', help='输入文档路径')
    parser.add_argument('--output', '-o', default='/workspace/data/quiz_questions.json', help='输出JSON文件路径')
    parser.add_argument('--title', '-t', default='知识测试', help='测试标题')
    parser.add_argument('--description', '-d', default='', help='测试描述')
    parser.add_argument('--num-questions', '-n', type=int, default=25, help='生成题目数量')
    parser.add_argument('--time-limit', '-l', type=int, default=30, help='答题时间限制（分钟）')
    parser.add_argument('--domain', help='知识领域 (medical/technical/business/legal)')
    
    args = parser.parse_args()
    
    # 检查文件是否存在
    if not os.path.exists(args.input_file):
        print(f"错误：文件不存在 {args.input_file}")
        return
    
    print(f"正在处理文档: {args.input_file}")
    
    # 根据文件扩展名选择提取方法
    file_ext = os.path.splitext(args.input_file)[1].lower()
    
    if file_ext == '.docx':
        content = extract_docx_content(args.input_file)
    elif file_ext == '.txt':
        content = extract_txt_content(args.input_file)
    elif file_ext == '.pdf':
        content = extract_pdf_content(args.input_file)
    else:
        print(f"不支持的文件格式: {file_ext}")
        print("支持的格式: .docx, .txt, .pdf")
        return
    
    if not content:
        print("无法提取文档内容")
        return
    
    print(f"文档内容长度: {len(content)} 字符")
    
    # 解析问答内容
    print("正在解析问答内容...")
    qa_pairs = parse_qa_content_universal(content)
    
    if not qa_pairs:
        print("错误：无法解析到有效的问答对")
        print("请检查文档格式是否符合要求")
        return
    
    print(f"成功解析得到 {len(qa_pairs)} 个问答对")
    
    # 显示前几个问答对预览
    print("\n=== 解析结果预览 ===")
    for i, qa in enumerate(qa_pairs[:3]):
        print(f"\n问答对 {i+1}:")
        print(f"问题: {qa['question']}")
        print(f"答案: {qa['answer'][:100]}...")
    
    # 生成题目配置
    config = {
        'title': args.title,
        'description': args.description or f"基于{os.path.basename(args.input_file)}生成的测试题目",
        'num_questions': args.num_questions,
        'time_limit': args.time_limit,
        'domain': args.domain or ''
    }
    
    # 生成题目
    print(f"\n正在生成 {min(args.num_questions, len(qa_pairs))} 道题目...")
    quiz_data = generate_quiz_questions(qa_pairs, config)
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(args.output), exist_ok=True)
    
    # 保存题目数据
    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(quiz_data, f, ensure_ascii=False, indent=2)
    
    print(f"题目生成完成！已保存到: {args.output}")
    
    # 显示题目预览
    print("\n=== 题目预览 ===")
    for i, q in enumerate(quiz_data['questions'][:3]):
        print(f"\n第 {q['id']} 题: {q['question']}")
        for j, option in enumerate(q['options']):
            marker = "★" if j == q['correct_answer'] else " "
            print(f"  {chr(65+j)}. {option[:60]}{'...' if len(option) > 60 else ''} {marker}")
    
    print(f"\n✅ 成功生成 {len(quiz_data['questions'])} 道题目")
    print(f"📝 题目标题: {quiz_data['title']}")
    print(f"⏱️ 答题时限: {quiz_data['time_limit']} 分钟")

if __name__ == "__main__":
    main()
