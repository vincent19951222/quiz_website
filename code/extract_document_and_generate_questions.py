#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
糖尿病知识文档内容提取和题目生成工具
"""

import json
import re
import random
from typing import List, Dict, Any
from docx import Document
import os

def extract_docx_content(file_path: str) -> str:
    """提取Word文档内容"""
    try:
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text.strip() + "\n"
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content += " | ".join(row_text) + "\n"
        
        return content
    except Exception as e:
        print(f"提取文档内容时出错: {e}")
        return ""

def parse_qa_content(content: str) -> List[Dict[str, str]]:
    """解析问答内容，提取问题和答案"""
    qa_pairs = []
    
    # 按行分割内容
    lines = content.split('\n')
    current_question = ""
    current_answer = ""
    in_answer = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 匹配问题模式（数字开头的问题）
        question_match = re.match(r'^\d+[、.]?\s*(.+)', line)
        if question_match:
            # 保存上一个问答对
            if current_question and current_answer:
                qa_pairs.append({
                    "question": current_question,
                    "answer": current_answer
                })
            
            current_question = question_match.group(1)
            current_answer = ""
            in_answer = False
            continue
        
        # 匹配答案模式
        if re.match(r'^(答|答案|解答)[：:]', line):
            in_answer = True
            answer_text = re.sub(r'^(答|答案|解答)[：:]\s*', '', line)
            current_answer = answer_text
            continue
        
        # 如果在答案模式中，继续添加答案内容
        if in_answer:
            current_answer += " " + line
        elif current_question:
            # 如果还没有开始答案，但有问题，可能是问题的续行
            current_question += " " + line
    
    # 保存最后一个问答对
    if current_question and current_answer:
        qa_pairs.append({
            "question": current_question,
            "answer": current_answer
        })
    
    return qa_pairs

def generate_wrong_options(correct_answer: str, context: str) -> List[str]:
    """基于正确答案和上下文生成错误选项"""
    wrong_options = []
    
    # 提取数字类的错误选项
    numbers = re.findall(r'\d+(?:\.\d+)?', correct_answer)
    if numbers:
        for num in numbers[:2]:  # 最多取前2个数字
            try:
                base_num = float(num)
                # 生成相近但错误的数字
                wrong_num1 = str(int(base_num * 1.2)) if base_num > 1 else str(round(base_num * 1.2, 1))
                wrong_num2 = str(int(base_num * 0.8)) if base_num > 1 else str(round(base_num * 0.8, 1))
                
                wrong_opt1 = correct_answer.replace(num, wrong_num1)
                wrong_opt2 = correct_answer.replace(num, wrong_num2)
                
                if wrong_opt1 != correct_answer:
                    wrong_options.append(wrong_opt1)
                if wrong_opt2 != correct_answer and len(wrong_options) < 2:
                    wrong_options.append(wrong_opt2)
            except:
                pass
    
    # 生成关键词替换的错误选项
    keywords_replacement = {
        '增加': '减少', '减少': '增加', '升高': '降低', '降低': '升高',
        '上升': '下降', '下降': '上升', '提高': '降低', '增强': '减弱',
        '不能': '能够', '禁止': '允许', '避免': '推荐', '预防': '治疗',
        '早期': '晚期', '急性': '慢性', '局部': '全身', '轻度': '重度',
        '正常': '异常', '健康': '患病', '有效': '无效', '安全': '危险'
    }
    
    for original, replacement in keywords_replacement.items():
        if original in correct_answer and len(wrong_options) < 3:
            wrong_option = correct_answer.replace(original, replacement)
            if wrong_option != correct_answer and wrong_option not in wrong_options:
                wrong_options.append(wrong_option)
    
    # 如果还不够3个选项，生成通用错误选项
    generic_wrong_options = [
        "以上都不正确",
        "需要更多信息才能确定",
        "取决于具体情况",
        "没有明确规定",
        "因人而异"
    ]
    
    for option in generic_wrong_options:
        if len(wrong_options) < 3 and option not in wrong_options:
            wrong_options.append(option)
    
    return wrong_options[:3]  # 确保只返回3个错误选项

def generate_multiple_choice_questions(qa_pairs: List[Dict[str, str]], num_questions: int = 25) -> List[Dict[str, Any]]:
    """基于问答对生成单选题"""
    questions = []
    
    # 如果问答对不够，使用全部；如果太多，随机选择
    if len(qa_pairs) <= num_questions:
        selected_pairs = qa_pairs
    else:
        selected_pairs = random.sample(qa_pairs, num_questions)
    
    for i, qa in enumerate(selected_pairs, 1):
        question_text = qa['question']
        correct_answer = qa['answer']
        
        # 生成错误选项
        wrong_options = generate_wrong_options(correct_answer, question_text)
        
        # 组合所有选项
        all_options = [correct_answer] + wrong_options
        random.shuffle(all_options)  # 随机排列选项
        
        # 找到正确答案的索引
        correct_index = all_options.index(correct_answer)
        
        question_data = {
            "id": i,
            "question": question_text,
            "options": all_options,
            "correct_answer": correct_index,
            "explanation": f"根据糖尿病基础知识：{correct_answer}"
        }
        
        questions.append(question_data)
    
    return questions

def main():
    """主函数"""
    # 文档路径
    doc_path = "/workspace/user_input_files/3.4糖尿病基础知识培训100问.docx"
    
    if not os.path.exists(doc_path):
        print(f"文档文件不存在: {doc_path}")
        return
    
    print("正在提取文档内容...")
    content = extract_docx_content(doc_path)
    
    if not content:
        print("无法提取文档内容")
        return
    
    print(f"文档内容长度: {len(content)} 字符")
    
    # 保存原始内容用于调试
    with open('/workspace/data/extracted_content.txt', 'w', encoding='utf-8') as f:
        f.write(content)
    print("原始内容已保存到 /workspace/data/extracted_content.txt")
    
    print("正在解析问答内容...")
    qa_pairs = parse_qa_content(content)
    
    print(f"解析得到 {len(qa_pairs)} 个问答对")
    
    if len(qa_pairs) < 10:
        print("警告：解析的问答对数量较少，可能需要调整解析逻辑")
        # 显示前几个问答对用于调试
        for i, qa in enumerate(qa_pairs[:5]):
            print(f"\n问答对 {i+1}:")
            print(f"问题: {qa['question'][:100]}...")
            print(f"答案: {qa['answer'][:100]}...")
    
    print("正在生成单选题...")
    questions = generate_multiple_choice_questions(qa_pairs, 25)
    
    # 保存生成的题目
    output_data = {
        "title": "糖尿病基础知识测试",
        "description": "基于糖尿病基础知识培训100问生成的测试题目",
        "time_limit": 30,  # 30分钟
        "questions": questions
    }
    
    with open('/workspace/data/quiz_questions.json', 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print(f"成功生成 {len(questions)} 道单选题")
    print("题目数据已保存到 /workspace/data/quiz_questions.json")
    
    # 显示前几道题目作为预览
    print("\n=== 题目预览 ===")
    for i, q in enumerate(questions[:3]):
        print(f"\n第 {q['id']} 题: {q['question']}")
        for j, option in enumerate(q['options']):
            marker = "★" if j == q['correct_answer'] else " "
            print(f"  {chr(65+j)}. {option} {marker}")
        print(f"解析: {q['explanation']}")

if __name__ == "__main__":
    # 确保数据目录存在
    os.makedirs('/workspace/data', exist_ok=True)
    main()
